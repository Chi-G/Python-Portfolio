from tkinter.tix import Tree
from turtle import title, width
from xml.dom.minidom import Document
from docx.shared import Inches

from docx import Document

document = Document()

document.add_picture(
    'chijindu.png', 
    width=Inches(0.8)
)

# Portfolio contact credentials
document.add_heading('Chijindu Nwokeohuru' + '\n').bold = True

home_address = input('input your home address: ')
city_country = input('input your city and country with a coma: ')
phone = input('input your phone number: ' )
email = input('input your email address: ')
portfolio = input('input your portfolio link: ')

document.add_paragraph(home_address + ' | ' + city_country + ' | ' + email + ' | ' + phone + ' | ' + portfolio)

# About me or Career Objectives credentials
document.add_heading('CAREER OBJECTIVES')
p = document.add_paragraph('******************************************************************************************* ')
document.add_paragraph(
    input('State your career objectives: ')
)

# Education
document.add_heading('EDUCATION')
p = document.add_paragraph('******************************************************************************************* ')
school = input('Enter name of institution: ')
course = input('Enter name or course of study: ')
start_date = input('Start Year: ')
end_date = input('End year: ')

p.add_run(school + ' ' + ' | ').bold = True
p.add_run(course + ' | ')
p.add_run(start_date + '-' + end_date + '\n').italic = True

course_details = input(
    'State your major area on ' + course + ' ' + 'while in' + ' ' + school + ': '
)
p.add_run(course_details)

# More education
while True:
    add_more_education = input(
        'Do you have additional education? y/n:')
    if add_more_education.lower() == 'y':
        school = input('Enter name of institution: ')
        course = input('Enter name or course of study: ')
        start_date = input('Start Year:')
        end_date = input('End year:')

        p.add_run(course +  ' ' + ' | ').bold = True
        p.add_run(course)
        p.add_run(start_date + '-' + end_date + '\n').italic = True

        course_details = input(
               'State your major area on ' + course + ' ' + 'while in' + ' ' + school + ': '
        )
        p.add_run(course_details)
    else:
        break

# Skills & qualifications
document.add_heading('Skills')
p = document.add_paragraph('*******************************************************************************************' )
skill = input('Enter your skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    add_more_skills = input('Do you have more skills? y/n')
    if add_more_skills.lower() == 'y':
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# work experience credentials
document.add_heading('Work Experience')
p = document.add_paragraph('*******************************************************************************************')
company = input('Enter your Company name: ')
job_title = input('Enter your job position: ')
start_date = input('Enter Start Year: ')
end_date = input('Enter End year: ')

p.add_run(company + ' ' ).bold = True
p.add_run(' | ' + start_date + '-' + end_date).italic = True

experience_details = input(
    'State your job description at ' + company + ': '
)
p.add_run(experience_details)

#more job description 1
while True:
    more_description =  input(
        'do you have more job description at' + ' ' + company + 'y/n' + ' ' +  '?:')
    if more_description.lower() == 'y' :
        experience_details = input(
            'State your job description at ' + company + ': ')
        p.add_run(experience_details)
        p.style = 'List Bullet'
    else:
        break

# More experience
while True:
    add_more_experience = input(
        'Do you have additional experience? y/n: ')
    if add_more_experience.lower() == 'y':
        company = input('Enter Company name:')
        job_title = input('Enter your job title: ')
        start_date = input('Start Year:')
        end_date = input('End year:')

        experience_details = input(
            'State your job description at ' + company + ':'
        )
        p.add_run(experience_details)
        p.style = 'List Bullet'

        #more job description 2
        while True:
                more_description =  input(
                    'do you have more job description at' + ' ' + company + 'y/n' + ' ' +  '?:')
                if more_description.lower() == 'y' :
                    experience_details = input(
                        'State your job description at ' + company + ': ')
                    p.add_run(experience_details)
                    p.style = 'List Bullet'
                else:
                    break

        p.add_run(company +  ' ' + ' | ').bold = True
        p.add_run(' | ' + start_date + '-' + end_date).italic = True

        p.add_run(experience_details)
        p.style = 'List Bullet'

    else:
        break


# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Python by Chijindu Nwokeohuru'

document.save('cv.docx')